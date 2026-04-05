"""
Multi-Format Document Loaders
==============================

Each loader converts a file into ``langchain_core.documents.Document``
objects with rich metadata.  Gemini Vision + Cloud Vision are used for
scanned PDFs and images; pdfplumber and pypdf serve as fallbacks.

Tabular files (Excel / CSV) produce *multiple* document representations
(full data, summary, row chunks, column statistics) so that both
list-all and aggregation queries return accurate results.
"""

from __future__ import annotations

import io
import json
import re
from abc import ABC, abstractmethod
from datetime import datetime
from pathlib import Path
from typing import Any, Dict, List, Optional

import numpy as np
import pandas as pd
from langchain_core.documents import Document

from config.settings import Settings, get_settings
from src.utils.helpers import clean_text, format_value, truncate
from src.utils.logger import get_logger

logger = get_logger(__name__)

# ── Cloud OCR (Gemini Vision + Cloud Vision) ──────────────────────────────

from src.ingestion.ocr_manager import (
    OCR_AVAILABLE as _OCR_AVAILABLE,
    get_gemini_ocr,
    get_cloud_vision,
    ocr_image,
)

try:
    from rank_bm25 import BM25Okapi  # noqa: F401 — verify install

    _BM25_AVAILABLE = True
except ImportError:
    _BM25_AVAILABLE = False


# ═══════════════════════════════════════════════════════════════════════════
#  BASE CLASS
# ═══════════════════════════════════════════════════════════════════════════


class BaseLoader(ABC):
    """Abstract base for all file loaders.

    Subclasses must implement ``load`` which returns a list of LangChain
    ``Document`` objects, each with ``page_content`` and ``metadata``.
    """

    @abstractmethod
    def load(
        self,
        file_path: Path,
        content: Optional[bytes] = None,
    ) -> List[Document]:
        """Load a file and return one or more ``Document`` objects.

        Args:
            file_path: Path to the file on disk **or** a virtual filename
                       when *content* bytes are supplied directly.
            content:   Raw file bytes (e.g. from an upload).  When given
                       the loader should read from these bytes instead of
                       opening *file_path*.

        Returns:
            A list of LangChain ``Document`` instances.
        """
        ...

    @staticmethod
    def _base_metadata(file_path: Path, content: Optional[bytes] = None) -> Dict[str, Any]:
        """Build base metadata common to every document."""
        meta: Dict[str, Any] = {
            "source": str(file_path),
            "filename": file_path.name,
            "file_extension": file_path.suffix.lower(),
        }
        if content is not None:
            meta["file_size_bytes"] = len(content)
        elif file_path.exists():
            stat = file_path.stat()
            meta["file_size_bytes"] = stat.st_size
            meta["modified_time"] = datetime.fromtimestamp(stat.st_mtime).isoformat()
        return meta


# ═══════════════════════════════════════════════════════════════════════════
#  PDF LOADER  — pdfplumber → Gemini Vision OCR → pypdf fallback chain
# ═══════════════════════════════════════════════════════════════════════════


class PDFLoader(BaseLoader):
    """Advanced PDF loader with Gemini Vision OCR for scanned pages.

    Extraction pipeline:
      1. **pdfplumber** — fast text + basic tables for digital PDFs.
      2. **Gemini Vision OCR** — cloud-based intelligent OCR for
         scanned/image-heavy pages.
      3. **pypdf** — last-resort plain-text extraction.

    When GOOGLE_API_KEY is not set the loader still works using
    pdfplumber and pypdf only.
    """

    def __init__(self) -> None:
        self._gemini = get_gemini_ocr()
        self._cloud = get_cloud_vision()
        if self._gemini:
            logger.info("pdf_loader_ocr_ready", backend="gemini_vision")
        elif self._cloud:
            logger.info("pdf_loader_ocr_ready", backend="cloud_vision")

    # ── public API ────────────────────────────────────────────────────

    def load(
        self,
        file_path: Path,
        content: Optional[bytes] = None,
    ) -> List[Document]:
        base_meta = self._base_metadata(file_path, content)
        source = io.BytesIO(content) if content else file_path

        documents: List[Document] = []

        try:
            pages_text, tables_text, page_count = self._extract_pdfplumber(source)
        except Exception as exc:
            logger.warning("pdfplumber_extraction_failed", error=str(exc))
            pages_text, tables_text, page_count = [], [], 0

        # If pdfplumber yielded very little text, try Gemini Vision OCR
        # (500 chars threshold — below this, it's almost certainly scanned)
        total_chars = sum(len(t) for t in pages_text)
        used_cloud_ocr = False
        if total_chars < 500 and (self._gemini or self._cloud):
            try:
                pdf_bytes = content if content else (
                    file_path.read_bytes() if file_path.exists() else b""
                )
                if pdf_bytes:
                    ocr_text, ocr_tables = self._extract_cloud_ocr(pdf_bytes)
                    if sum(len(t) for t in ocr_text) > total_chars:
                        pages_text = ocr_text
                        tables_text = ocr_tables
                        base_meta["extraction_method"] = "gemini_vision"
                        used_cloud_ocr = True
                        if page_count == 0:
                            page_count = len(ocr_text)
            except Exception as exc:
                logger.warning("cloud_ocr_extraction_failed", error=str(exc))

        # Last fallback: pypdf (rewind BytesIO first!)
        if not pages_text:
            try:
                if isinstance(source, io.BytesIO):
                    source.seek(0)
                pages_text, page_count = self._extract_pypdf(source)
                base_meta["extraction_method"] = "pypdf"
            except Exception as exc:
                logger.error("all_pdf_extraction_failed", error=str(exc))
                return []

        # ALWAYS extract embedded images and OCR them, even if we got text
        # (common case: PDFs with text + embedded figures/charts/diagrams)
        if not used_cloud_ocr and (self._gemini or self._cloud):
            try:
                pdf_bytes = content if content else (
                    file_path.read_bytes() if file_path.exists() else b""
                )
                if pdf_bytes:
                    pages_text = self._extract_embedded_images(
                        pdf_bytes, pages_text
                    )
            except Exception as exc:
                logger.debug("embedded_image_extraction_failed", error=str(exc))

        base_meta.setdefault("extraction_method", "pdfplumber")
        base_meta["page_count"] = page_count
        base_meta["file_type"] = "pdf"

        # Build full-document Document
        full_content = "\n\n---\n\n".join(
            t for t in pages_text if t and t.strip()
        )
        if full_content.strip():
            documents.append(
                Document(
                    page_content=f"# PDF: {file_path.name}\nPages: {page_count}\n\n{full_content}",
                    metadata={**base_meta, "document_type": "full_data", "priority": "high"},
                )
            )

        # Per-page Documents
        for page_num, text in enumerate(pages_text, 1):
            if not text or not text.strip():
                continue
            page_meta = {
                **base_meta,
                "document_type": "page",
                "page_number": page_num,
                "total_pages": page_count,
            }
            page_text = text
            if page_num <= len(tables_text) and tables_text[page_num - 1]:
                page_text += "\n\n" + tables_text[page_num - 1]
            documents.append(
                Document(
                    page_content=f"[PDF: {file_path.name} | Page {page_num}]\n\n{page_text}",
                    metadata=page_meta,
                )
            )

        logger.info(
            "pdf_loaded",
            filename=file_path.name,
            pages=page_count,
            documents=len(documents),
            method=base_meta.get("extraction_method"),
        )
        return documents

    # ── extraction methods ────────────────────────────────────────────

    def _extract_pdfplumber(self, source: Any):
        """Extract text + tables with pdfplumber."""
        import pdfplumber

        pages_text: List[str] = []
        tables_text: List[str] = []
        page_count = 0

        with pdfplumber.open(source) as pdf:
            page_count = len(pdf.pages)
            for page in pdf.pages:
                text = page.extract_text() or ""
                pages_text.append(clean_text(text))

                # Extract tables as formatted text
                page_tables = page.extract_tables()
                tbl_parts: List[str] = []
                if page_tables:
                    for i, table in enumerate(page_tables, 1):
                        tbl_parts.append(f"[Table {i}]")
                        for row in table:
                            row_str = " | ".join(
                                str(c) if c else "" for c in row
                            )
                            tbl_parts.append(row_str)
                tables_text.append("\n".join(tbl_parts))

        return pages_text, tables_text, page_count

    def _extract_cloud_ocr(self, pdf_bytes: bytes):
        """Extract text + tables with Gemini Vision / Cloud Vision OCR.

        Renders each PDF page to an image using PyMuPDF, then sends it
        to the cloud OCR backends for extraction.
        """
        import fitz  # PyMuPDF

        doc_pdf = fitz.open("pdf", pdf_bytes)
        pages_text: List[str] = []
        tables_text: List[str] = []

        for page_obj in doc_pdf:
            # Render page at 300 DPI
            pix = page_obj.get_pixmap(dpi=300)
            img_np = np.frombuffer(pix.samples, dtype=np.uint8).reshape(
                pix.height, pix.width, pix.n
            )
            if img_np.shape[2] == 4:
                img_np = img_np[:, :, :3]  # drop alpha
            elif img_np.shape[2] == 1:
                img_np = np.stack([img_np.squeeze()] * 3, axis=-1)

            # Run cloud OCR
            text, conf = ocr_image(img_np)
            pages_text.append(clean_text(text))
            tables_text.append("")  # Tables are embedded in the OCR text

        doc_pdf.close()
        return pages_text, tables_text

    def _extract_embedded_images(
        self, pdf_bytes: bytes, pages_text: List[str]
    ) -> List[str]:
        """Extract embedded images from a PDF and OCR them.

        This handles the common case where the PDF has digital text
        PLUS embedded figures, charts, or diagrams with text in them.
        Each extracted image is OCR'd and appended to the corresponding page.
        """
        import fitz  # PyMuPDF

        # Copy the list so we don't mutate the caller's data
        result = list(pages_text)

        try:
            doc_pdf = fitz.open("pdf", pdf_bytes)
            for page_num, page_obj in enumerate(doc_pdf):
                image_list = page_obj.get_images(full=True)
                for img_info in image_list:
                    xref = img_info[0]
                    try:
                        base_image = doc_pdf.extract_image(xref)
                        if not base_image or not base_image.get("image"):
                            continue
                        image_bytes = base_image["image"]
                        from PIL import Image as PILImage
                        pil_img = PILImage.open(io.BytesIO(image_bytes)).convert("RGB")
                        w, h = pil_img.size
                        # Skip tiny images (icons, bullets, decorations)
                        if w < 80 or h < 80:
                            continue
                        # Skip very small area images (< ~100x100)
                        if w * h < 10000:
                            continue
                        img_arr = np.array(pil_img)
                        img_text, _ = ocr_image(img_arr)
                        if img_text.strip() and len(img_text.strip()) > 15:
                            if page_num < len(result):
                                result[page_num] += f"\n\n[Embedded Image OCR]:\n{img_text}"
                            else:
                                result.append(f"[Embedded Image OCR]:\n{img_text}")
                            logger.debug(
                                "embedded_image_ocr",
                                page=page_num + 1,
                                size=f"{w}x{h}",
                                chars=len(img_text),
                            )
                    except Exception as e:
                        logger.debug(
                            "embedded_image_extraction_error",
                            page=page_num + 1,
                            error=str(e),
                        )
                        continue
            doc_pdf.close()
        except Exception as exc:
            logger.debug("embedded_image_pass_failed", error=str(exc))

        return result

    def _extract_pypdf(self, source: Any):
        """Last-resort extraction with pypdf."""
        from pypdf import PdfReader

        if isinstance(source, io.BytesIO):
            source.seek(0)

        reader = PdfReader(source)
        pages_text = []
        for page in reader.pages:
            text = page.extract_text() or ""
            pages_text.append(clean_text(text))
        return pages_text, len(reader.pages)

    @staticmethod
    def _html_table_to_text(html: str) -> str:
        """Convert HTML table string to pipe-delimited text."""
        try:
            dfs = pd.read_html(io.StringIO(html))
            if dfs:
                return dfs[0].to_string(index=False)
        except Exception:
            pass
        text = re.sub(r"<[^>]+>", " ", html)
        return re.sub(r"\s+", " ", text).strip()


# ═══════════════════════════════════════════════════════════════════════════
#  DOCX LOADER
# ═══════════════════════════════════════════════════════════════════════════


class DocxLoader(BaseLoader):
    """Word document loader — paragraphs + tables."""

    def load(self, file_path: Path, content: Optional[bytes] = None) -> List[Document]:
        from docx import Document as DocxDocument

        base_meta = self._base_metadata(file_path, content)
        base_meta["file_type"] = "docx"

        try:
            if content:
                doc = DocxDocument(io.BytesIO(content))
            else:
                doc = DocxDocument(str(file_path))
        except Exception as exc:
            logger.error("docx_load_failed", filename=file_path.name, error=str(exc))
            return []

        parts: List[str] = [f"# Word Document: {file_path.name}\n"]

        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            style = para.style.name if para.style else "Normal"
            if "Heading" in style:
                parts.append(f"\n## {text}")
            else:
                parts.append(text)

        for i, table in enumerate(doc.tables, 1):
            parts.append(f"\n### Table {i}")
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells)
                parts.append(row_text)

        # OCR embedded images using cloud OCR
        try:
            for rel in doc.part.rels.values():
                if "image" in getattr(rel, "reltype", ""):
                    blob = rel.target_part.blob
                    from PIL import Image as PILImage
                    pil = PILImage.open(io.BytesIO(blob)).convert("RGB")
                    arr = np.array(pil)
                    text, _ = ocr_image(arr)
                    if text.strip():
                        parts.append(f"\n[Embedded Image]: {text}")
        except Exception:
            pass

        full_text = "\n".join(parts)
        documents = [
            Document(
                page_content=full_text,
                metadata={**base_meta, "document_type": "full_data"},
            )
        ]
        logger.info("docx_loaded", filename=file_path.name, documents=len(documents))
        return documents


# ═══════════════════════════════════════════════════════════════════════════
#  EXCEL LOADER — multi-representation (full + summary + rows + columns)
# ═══════════════════════════════════════════════════════════════════════════


class ExcelLoader(BaseLoader):
    """Enhanced Excel loader producing multiple document representations.

    For every sheet the loader produces:
      1. **full_data** — every row serialised, so list-all queries work.
      2. **summary** — column statistics (sum, mean, unique counts).
      3. **rows** — row-level chunks (50 rows each) for efficient retrieval.
      4. **columns** — per-column value lists for aggregation queries.
    """

    MAX_ROWS_PER_CHUNK: int = 50

    def load(self, file_path: Path, content: Optional[bytes] = None) -> List[Document]:
        base_meta = self._base_metadata(file_path, content)
        base_meta["file_type"] = "excel"
        documents: List[Document] = []

        try:
            engine = "openpyxl" if file_path.suffix.lower() == ".xlsx" else "xlrd"
            src = io.BytesIO(content) if content else file_path
            xls = pd.ExcelFile(src, engine=engine)
        except Exception as exc:
            logger.error("excel_load_failed", filename=file_path.name, error=str(exc))
            return []

        for sheet in xls.sheet_names:
            df = pd.read_excel(xls, sheet_name=sheet)
            if df.empty:
                continue
            df.columns = df.columns.astype(str).str.strip()

            documents.append(self._full_data_doc(file_path.name, sheet, df, base_meta))
            documents.append(self._summary_doc(file_path.name, sheet, df, base_meta))
            documents.extend(self._row_docs(file_path.name, sheet, df, base_meta))
            documents.extend(self._column_docs(file_path.name, sheet, df, base_meta))

        logger.info("excel_loaded", filename=file_path.name, documents=len(documents))
        return documents

    # ── representation builders ───────────────────────────────────────

    def _full_data_doc(
        self, fname: str, sheet: str, df: pd.DataFrame, base: Dict
    ) -> Document:
        lines = [
            f"# COMPLETE DATA: {fname}",
            f"## Sheet: {sheet}",
            f"## Total Records: {len(df)}",
            f"## Columns: {', '.join(df.columns)}",
            "",
            "## ALL DATA ROWS:",
            "",
        ]
        for idx, row in df.iterrows():
            parts = [f"{col}: {format_value(row[col])}" for col in df.columns]
            lines.append(f"Row {idx + 1}: {' | '.join(parts)}")
        return Document(
            page_content="\n".join(lines),
            metadata={
                **base,
                "document_type": "full_data",
                "sheet_name": sheet,
                "total_rows": len(df),
                "columns": list(df.columns),
                "priority": "high",
            },
        )

    def _summary_doc(
        self, fname: str, sheet: str, df: pd.DataFrame, base: Dict
    ) -> Document:
        lines = [
            f"# DATA SUMMARY: {fname} - {sheet}",
            "",
            f"- Total Rows: {len(df)}",
            f"- Total Columns: {len(df.columns)}",
            f"- Column Names: {', '.join(df.columns)}",
            "",
        ]
        for col in df.select_dtypes(include=["number"]).columns:
            data = df[col].dropna()
            if data.empty:
                continue
            lines.extend([
                f"\n### {col}",
                f"- Count: {len(data)}",
                f"- Sum: {data.sum():,.2f}",
                f"- Mean: {data.mean():,.2f}",
                f"- Min: {data.min():,.2f}",
                f"- Max: {data.max():,.2f}",
            ])
        for col in df.select_dtypes(include=["object"]).columns[:5]:
            vc = df[col].value_counts().head(5)
            lines.append(f"\n### {col} — {df[col].nunique()} unique")
            for val, cnt in vc.items():
                lines.append(f"  - {val}: {cnt}")
        return Document(
            page_content="\n".join(lines),
            metadata={**base, "document_type": "summary", "sheet_name": sheet},
        )

    def _row_docs(
        self, fname: str, sheet: str, df: pd.DataFrame, base: Dict
    ) -> List[Document]:
        docs: List[Document] = []
        total = len(df)
        for start in range(0, total, self.MAX_ROWS_PER_CHUNK):
            end = min(start + self.MAX_ROWS_PER_CHUNK, total)
            chunk_df = df.iloc[start:end]
            lines = [f"# DATA ROWS: {fname} - {sheet}", f"## Rows {start + 1}–{end} of {total}", ""]
            for idx, row in chunk_df.iterrows():
                parts = [f"{c}: {format_value(row[c])}" for c in df.columns]
                lines.append(f"Row {idx + 1}: {' | '.join(parts)}")
            docs.append(
                Document(
                    page_content="\n".join(lines),
                    metadata={
                        **base,
                        "document_type": "rows",
                        "sheet_name": sheet,
                        "row_range": f"{start + 1}-{end}",
                        "total_rows": total,
                    },
                )
            )
        return docs

    def _column_docs(
        self, fname: str, sheet: str, df: pd.DataFrame, base: Dict
    ) -> List[Document]:
        docs: List[Document] = []
        for col in df.columns:
            data = df[col].dropna()
            if data.empty:
                continue
            lines = [f"# COLUMN DATA: {col}", f"Source: {fname} - {sheet}", f"Total Values: {len(data)}", ""]
            if pd.api.types.is_numeric_dtype(data):
                lines.extend([
                    f"Sum: {data.sum():,.2f}",
                    f"Mean: {data.mean():,.2f}",
                    f"Min: {data.min():,.2f}",
                    f"Max: {data.max():,.2f}",
                    "",
                    "## All Values:",
                ])
                for idx, val in data.items():
                    lines.append(f"Row {idx + 1}: {format_value(val)}")
            else:
                vc = data.value_counts()
                lines.append(f"Unique Values: {data.nunique()}")
                for val, cnt in vc.items():
                    lines.append(f"- {val}: {cnt}")
            docs.append(
                Document(
                    page_content="\n".join(lines),
                    metadata={**base, "document_type": "column", "sheet_name": sheet, "column_name": col},
                )
            )
        return docs


# ═══════════════════════════════════════════════════════════════════════════
#  CSV LOADER — reuses Excel logic
# ═══════════════════════════════════════════════════════════════════════════


class CSVLoader(BaseLoader):
    """CSV loader — delegates to ``ExcelLoader`` representations."""

    ENCODINGS = ("utf-8", "latin-1", "cp1252", "iso-8859-1", "utf-16")

    def load(self, file_path: Path, content: Optional[bytes] = None) -> List[Document]:
        df = self._read(file_path, content)
        if df is None or df.empty:
            logger.warning("csv_empty", filename=file_path.name)
            return []

        base_meta = self._base_metadata(file_path, content)
        base_meta["file_type"] = "csv"
        df.columns = df.columns.astype(str).str.strip()

        excel = ExcelLoader()
        fname = file_path.name
        docs: List[Document] = [
            excel._full_data_doc(fname, "CSV", df, base_meta),
            excel._summary_doc(fname, "CSV", df, base_meta),
            *excel._row_docs(fname, "CSV", df, base_meta),
            *excel._column_docs(fname, "CSV", df, base_meta),
        ]
        for d in docs:
            d.metadata["file_type"] = "csv"

        logger.info("csv_loaded", filename=file_path.name, documents=len(docs))
        return docs

    def _read(self, path: Path, content: Optional[bytes]) -> Optional[pd.DataFrame]:
        for enc in self.ENCODINGS:
            try:
                src = io.BytesIO(content) if content else path
                return pd.read_csv(src, encoding=enc)
            except Exception:
                if content:
                    continue
                continue
        return None


# ═══════════════════════════════════════════════════════════════════════════
#  TEXT / MARKDOWN LOADER
# ═══════════════════════════════════════════════════════════════════════════


class TextLoader(BaseLoader):
    """Plain text and Markdown loader."""

    def load(self, file_path: Path, content: Optional[bytes] = None) -> List[Document]:
        base_meta = self._base_metadata(file_path, content)
        ext = file_path.suffix.lower()
        base_meta["file_type"] = "markdown" if ext == ".md" else "text"

        text: Optional[str] = None
        raw = content if content else None

        for enc in ("utf-8", "latin-1", "cp1252"):
            try:
                if raw is not None:
                    text = raw.decode(enc)
                else:
                    text = file_path.read_text(encoding=enc)
                break
            except Exception:
                continue

        if not text:
            logger.warning("text_load_failed", filename=file_path.name)
            return []

        text = clean_text(text)
        doc = Document(
            page_content=f"# {base_meta['file_type'].title()}: {file_path.name}\n\n{text}",
            metadata={**base_meta, "document_type": "full_data", "word_count": len(text.split())},
        )
        logger.info("text_loaded", filename=file_path.name)
        return [doc]


# ═══════════════════════════════════════════════════════════════════════════
#  JSON LOADER
# ═══════════════════════════════════════════════════════════════════════════


class JSONLoader(BaseLoader):
    """JSON loader — arrays are split into per-item documents."""

    def load(self, file_path: Path, content: Optional[bytes] = None) -> List[Document]:
        base_meta = self._base_metadata(file_path, content)
        base_meta["file_type"] = "json"

        try:
            raw = content.decode("utf-8") if content else file_path.read_text("utf-8")
            data = json.loads(raw)
        except Exception as exc:
            logger.error("json_load_failed", filename=file_path.name, error=str(exc))
            return []

        docs: List[Document] = []
        full = f"# JSON: {file_path.name}\n\n{json.dumps(data, indent=2, default=str, ensure_ascii=False)}"
        docs.append(
            Document(
                page_content=full,
                metadata={**base_meta, "document_type": "full_data", "json_type": type(data).__name__},
            )
        )

        if isinstance(data, list):
            for i, item in enumerate(data):
                docs.append(
                    Document(
                        page_content=f"[JSON Item {i + 1}]\n{json.dumps(item, indent=2, default=str)}",
                        metadata={**base_meta, "document_type": "array_item", "item_index": i},
                    )
                )

        logger.info("json_loaded", filename=file_path.name, documents=len(docs))
        return docs


# ═══════════════════════════════════════════════════════════════════════════
#  IMAGE LOADER — standalone image files via OCR
# ═══════════════════════════════════════════════════════════════════════════


class ImageLoader(BaseLoader):
    """Loader for standalone image files (PNG, JPG, GIF, WebP, BMP, TIFF).

    Uses the cloud OCR pipeline (Gemini Vision / Cloud Vision) to extract
    text from images.  Supports adaptive preprocessing for low-quality scans.
    """

    def load(self, file_path: Path, content: Optional[bytes] = None) -> List[Document]:
        base_meta = self._base_metadata(file_path, content)
        base_meta["file_type"] = "image"

        raw = content if content else file_path.read_bytes()

        try:
            from PIL import Image as PILImage

            pil = PILImage.open(io.BytesIO(raw)).convert("RGB")
            img_arr = np.array(pil)
        except Exception as exc:
            logger.error("image_load_failed", filename=file_path.name, error=str(exc))
            return []

        w, h = pil.size
        base_meta["image_size"] = f"{w}x{h}"

        # Run OCR through the full Gemini Vision / Cloud Vision pipeline
        text, confidence = ocr_image(img_arr)

        if not text or not text.strip():
            logger.warning("image_ocr_empty", filename=file_path.name)
            return []

        base_meta["extraction_method"] = "ocr"
        base_meta["ocr_confidence"] = round(confidence, 3)

        doc = Document(
            page_content=f"# Image: {file_path.name}\nSize: {w}x{h}\n\n{clean_text(text)}",
            metadata={**base_meta, "document_type": "full_data", "priority": "high"},
        )

        logger.info(
            "image_loaded",
            filename=file_path.name,
            size=f"{w}x{h}",
            chars=len(text),
            confidence=round(confidence, 3),
        )
        return [doc]


# ═══════════════════════════════════════════════════════════════════════════
#  FACTORY
# ═══════════════════════════════════════════════════════════════════════════


class LoaderFactory:
    """Resolve and invoke the correct loader for a given file type."""

    _LOADERS: Dict[str, BaseLoader] = {}

    @classmethod
    def _ensure_loaders(cls) -> None:
        if not cls._LOADERS:
            cls._LOADERS = {
                "pdf": PDFLoader(),
                "docx": DocxLoader(),
                "excel": ExcelLoader(),
                "csv": CSVLoader(),
                "text": TextLoader(),
                "markdown": TextLoader(),
                "json": JSONLoader(),
                "image": ImageLoader(),
            }

    @classmethod
    def load_file(
        cls,
        file_path: Path,
        content: Optional[bytes] = None,
    ) -> List[Document]:
        """Load a file, selecting the correct loader automatically.

        Args:
            file_path: Path or virtual filename.
            content: Raw bytes (from upload).

        Returns:
            List of ``Document`` objects (may be empty on error).
        """
        cls._ensure_loaders()
        settings = get_settings()
        ext = file_path.suffix.lower()
        file_type = settings.SUPPORTED_EXTENSIONS.get(ext)

        if not file_type:
            logger.warning("unsupported_extension", ext=ext)
            return []

        loader = cls._LOADERS.get(file_type)
        if loader is None:
            logger.warning("no_loader", file_type=file_type)
            return []

        try:
            return loader.load(file_path, content)
        except Exception as exc:
            logger.error("load_failed", filename=file_path.name, error=str(exc))
            return []