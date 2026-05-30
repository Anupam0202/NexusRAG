"""
DOCX loader resource guard tests.
"""

from __future__ import annotations

import io
from pathlib import Path

from config.settings import get_settings
from src.ingestion.loader import DocxLoader


def _docx_with_image_bytes() -> bytes:
    from docx import Document as DocxDocument
    from PIL import Image

    image_buffer = io.BytesIO()
    Image.new("RGB", (32, 32), color="white").save(image_buffer, format="PNG")
    image_buffer.seek(0)

    doc = DocxDocument()
    doc.add_paragraph("Invoice total is 1234 rupees.")
    doc.add_picture(image_buffer)

    output = io.BytesIO()
    doc.save(output)
    return output.getvalue()


def test_docx_embedded_image_ocr_skips_when_memory_constrained(monkeypatch):
    monkeypatch.setenv("CONSTRAINED_MEMORY", "true")
    monkeypatch.setenv("ENABLE_DOCX_EMBEDDED_IMAGE_OCR", "true")
    get_settings.cache_clear()

    def fail_ocr(*_args, **_kwargs):
        raise AssertionError("DOCX embedded OCR should be skipped")

    monkeypatch.setattr("src.ingestion.loader.ocr_image", fail_ocr)
    try:
        docs = DocxLoader().load(Path("invoice.docx"), _docx_with_image_bytes())
    finally:
        get_settings.cache_clear()

    assert len(docs) == 1
    assert "Invoice total is 1234 rupees." in docs[0].page_content
